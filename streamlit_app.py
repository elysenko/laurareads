import streamlit as st
import os
import json
import streamlit as st
import dropbox
from dotenv import load_dotenv, set_key, find_dotenv
from os.path import join, dirname
from io import BytesIO
from docx import Document
import string
import requests

try:
    dotenv_path = join(dirname(__file__), '.env')
    load_dotenv(dotenv_path,verbose=True, override=True)
except:
    load_dotenv(verbose=True, override=True)

# Global Variables
st.set_page_config(layout="wide")

if not 'key_incr' in st.session_state.keys():
    st.session_state.key_num = 0

if not 'disp_text' in st.session_state.keys():
    st.session_state.disp_text = ''
    
if not 'doc' in st.session_state:
    st.session_state.doc = None
    
if not 'filename' in st.session_state:
    st.session_state.filename = ''

# Functions
def list_files_in_dropbox(dbx,path='',list_files=[]):

    # Start at the root folder ('')
    
    try:
        # List the folder contents
        result = dbx.files_list_folder(path)
        
        # Loop through the entries (files and folders)
        while True:
            for entry in result.entries:
                if isinstance(entry, dropbox.files.FileMetadata):
                    list_files.append(entry.path_display)
                elif isinstance(entry, dropbox.files.FolderMetadata):
                    print(f"Entering folder: {entry.path_display}")
                    # Recursive call to list files in this subfolder
                    sub_path = os.path.join(path,entry.path_lower)
                    list_files = list_files_in_dropbox(dbx,sub_path,list_files)

            # Check if there's more data to retrieve (pagination)
            if result.has_more:
                result = dbx.files_list_folder_continue(result.cursor)
            else:
                break

    except dropbox.exceptions.ApiError as err:
        print(f"API error: {err}")
    
    return list_files
        
def dropbox_client():
    """Creates a dropbox client"""
    
    try:
        app_key = os.getenv("DROPBOX_APP_KEY")
        app_secret = os.getenv("DROPBOX_APP_SECRET")
        refresh_token = os.getenv("DROPBOX_REFRESH_TOKEN")
        dbx = dropbox.Dropbox(app_key = app_key,
                          app_secret = app_secret,
                          oauth2_refresh_token = refresh_token
                          )
        status = dbx.users_get_current_account()
    except:

        new_api_key = get_access_token()
        dbx = dropbox.Dropbox(new_api_key)
        status = dbx.users_get_current_account()

    return dbx

def get_refresh_token():
    """Gets the dropbox refresh token"""
    
    app_key = os.getenv("DROPBOX_APP_KEY")
    app_secret = os.getenv("DROPBOX_APP_SECRET")

    refresh_token='b2ys6W_AO5UAAAAAAAAnED7hk0bnmvgXEwJhu6MxMQY'.strip()
    
    # build the authorization URL:
    authorization_url = "https://www.dropbox.com/oauth2/authorize?client_id=%s&token_access_type=offline&response_type=code" % app_key
    
    # send the user to the authorization URL:
    print('Go to the following URL and allow access:')
    print(authorization_url)
    
    set_key(find_dotenv(), "DROPBOX_APP_KEY", app_key)
    set_key(find_dotenv(), "DROPBOX_APP_SECRET", app_secret)
    set_key(find_dotenv(), "DROPBOX_REFRESH_TOKEN", refresh_token)
    
    moe = 1
    
    return refresh_token

def get_access_token():
    # get the authorization code from the user:
    # authorization_code = raw_input('Enter the code:\n')
    app_key = os.getenv("DROPBOX_APP_KEY")
    app_secret = os.getenv("DROPBOX_APP_SECRET")
    refresh_token=os.getenv("DROPBOX_REFRESH_TOKEN")
    
    # exchange the authorization code for an access token:
    token_url = "https://api.dropboxapi.com/oauth2/token"
    params = {
        "grant_type": "authorization_code",
        "code": refresh_token,
        "client_id": app_key,
        "client_secret": app_secret,
    }
    r = requests.post(token_url, data=params)
    print(r.text)

    if r.status_code == 200:
        text_dic = json.loads(r.text)
        new_access_token = text_dic['access_token']
        new_refresh_token = text_dic['refresh_token']
        
        set_key(find_dotenv(), "DROPBOX_ACCESS_TOKEN", new_access_token)
        set_key(find_dotenv(), "DROPBOX_REFRESH_TOKEN", new_refresh_token)
        
        return new_access_token
    else:
        print("cannot get new access_totken")

def read_file_from_dropbox(dbx,file_path):
    # Create a Dropbox client using the access token
    try:
        # Download the file
        metadata, response = dbx.files_download(file_path)

        # For example, if this was a binary file (like an image or PDF), you could save it to disk:
        with BytesIO(response.content) as stream:
            try:
                doc = read_docx_from_bytesio(stream)
                return doc 
            except:
                text_content = 'CANNOT OPEN FILE...'    
        return text_content

    except dropbox.exceptions.ApiError as err:
        print(f"Error reading file from Dropbox: {err}")
        return None
    
def read_docx_from_bytesio(file_like_object):
    doc = Document(file_like_object)  # Use the BytesIO object
    
    return doc 
def display_text(doc=None,filename=''):
    
    if not doc is None:
        # Extract headings to create a TOC
        toc = []
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                toc.append((para.style.name, para.text))
        
        # Display TOC in Streamlit
        st.sidebar.title('Table of Contents')
        for style, heading in toc:
            heading = heading.translate(str.maketrans('', '', string.punctuation)).replace('  ',' ').strip()
            st.sidebar.markdown(f"- [{heading}](#{heading.replace(' ', '-').lower()})")
        
        # Display the document content
        all_text = '/n'.join([para.text for para in doc.paragraphs])
        if not st.session_state.disp_text == all_text:
        
            st.title(filename.split('.')[0])
            i = len(toc)
            
            for para in doc.paragraphs:
                if i < len(toc) -10:
                    i += 1
                else:
                    if para.style.name.startswith('Heading'):
                        heading = para.text
                        heading = heading.translate(str.maketrans('', '', string.punctuation)).replace('  ',' ').lower().strip()
                        st.markdown(f"## {heading}")
                    elif para.style.name.startswith("Folder:"):
                        pass
                    else:
                        st.write(para.text)
            
            st.session_state.disp_text = ''

def button_swap(full_path):
    st.session_state[full_path] = not st.session_state[full_path]
    
    return

# Recursive function to display the folder tree
def display_tree(tree, current_path=''):
    for key, value in sorted(tree.items()):
        # Create full path for current item
        full_path = f"{current_path}/{key}" if current_path else key
        path_key = f'{full_path}_{st.session_state.key_num}'
        
        
        if value is None:  # It's a file
            if st.button(f"📄 {key}", key=path_key):
                st.session_state.doc = handle_file_click(full_path)
                
        else:  # It's a folder
            if full_path not in st.session_state:
                st.session_state[full_path] = False
                
            # Display the button with the current expansion state
            st.button(f"📁 {key} {'🔽' if st.session_state[full_path] else '▶️'}", on_click=button_swap,args=(full_path,))
            
            # If the folder is expanded, recursively display its contents
            if st.session_state[full_path]:
                display_tree(value, full_path)
            
# Callback function when a file is clicked
def handle_file_click(file_path):
    dbx = dropbox_client()
    print('file_path: (next)')
    print(file_path)
    st.session_state.filename = file_path.split('/')[-1]
    full_file_path = '/' + file_path
    
    doc = read_file_from_dropbox(dbx,full_file_path)

    return doc

# Function to convert flat structure into a tree
def convert_to_tree(paths):
    tree = {}
    for path in paths:
        parts = path.split('/')
        current = tree
        for index, part in enumerate(parts):
            if part not in current:
                current[part] = None if index == len(parts) - 1 else {}
            current = current[part]
    
    return tree

# Example usage
if __name__=='__main__':

    col1,col2,col3 = st.columns([1,2,1])
    with col2:
        st.title("🎈 Laura's Reading List...")
    
    dbx = dropbox_client()
    list_files = list_files_in_dropbox(dbx)
    list_files = [file[1:] for file in list_files]
    list_files = [file.replace('File: ','') for file in list_files]
    tree = convert_to_tree(list_files)
    
    col1,col2,col3 = st.columns([3,1,8])
    
    with col1:
        
        display_tree(tree)
    with col3:
        display_text(st.session_state.doc,st.session_state.filename)



